# Copyright (c) 2025 Bytedance Ltd. and/or its affiliates
# SPDX-License-Identifier: MIT

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

from src.config.loader import get_int_env
from src.crawler.readability_extractor import ReadabilityExtractor
from src.rag.builder import build_retriever
from src.rag.ingestion_types import IngestedChunk, IngestedDocument
from src.rag.retriever import Retriever

logger = logging.getLogger(__name__)


@dataclass
class LoadedSection:
    text: str
    metadata: dict


@dataclass
class LoadedDocument:
    title: str
    sections: list[LoadedSection]


@dataclass
class IngestionSummary:
    files_seen: int = 0
    files_ingested: int = 0
    chunks_ingested: int = 0
    failures: list[str] | None = None

    def to_dict(self) -> dict:
        return {
            "files_seen": self.files_seen,
            "files_ingested": self.files_ingested,
            "chunks_ingested": self.chunks_ingested,
            "failures": self.failures or [],
        }


DEFAULT_EXTENSIONS = [
    ".md",
    ".markdown",
    ".txt",
    ".html",
    ".htm",
    ".pdf",
    ".docx",
]


def ingest_paths(
    paths: Iterable[str],
    chunk_size: int,
    chunk_overlap: int,
    extensions: Iterable[str] | None = None,
    root: str | None = None,
    dry_run: bool = False,
) -> IngestionSummary:
    summary = IngestionSummary(failures=[])
    files = list(_discover_files(paths, extensions))
    summary.files_seen = len(files)

    if not files:
        logger.warning("No files found for ingestion.")
        return summary

    retriever = build_retriever()
    if not retriever:
        raise RuntimeError("No RAG provider configured for ingestion.")
    if retriever.__class__.ingest_documents is Retriever.ingest_documents:
        raise RuntimeError("Selected RAG provider does not support ingestion.")

    for file_path in files:
        try:
            loaded = _load_document(file_path)
            doc = _build_ingested_document(
                file_path=file_path,
                loaded=loaded,
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                root=root,
            )
            if not doc.chunks:
                logger.warning("Skipping empty document: %s", file_path)
                continue
            if not dry_run:
                retriever.ingest_documents([doc])
            summary.files_ingested += 1
            summary.chunks_ingested += len(doc.chunks)
        except Exception as exc:
            logger.warning("Failed to ingest %s: %s", file_path, exc)
            summary.failures.append(f"{file_path}: {exc}")

    return summary


def _discover_files(
    paths: Iterable[str],
    extensions: Iterable[str] | None,
) -> Iterable[Path]:
    allowed = _normalize_extensions(extensions)
    for raw_path in paths:
        path = Path(raw_path)
        if path.is_file():
            if not allowed or path.suffix.lower() in allowed:
                yield path
            continue
        if not path.exists():
            continue
        for candidate in path.rglob("*"):
            if candidate.is_file():
                if not allowed or candidate.suffix.lower() in allowed:
                    yield candidate


def _normalize_extensions(extensions: Iterable[str] | None) -> set[str]:
    if not extensions:
        return set(ext.lower() for ext in DEFAULT_EXTENSIONS)
    normalized = set()
    for ext in extensions:
        ext = ext.strip().lower()
        if not ext:
            continue
        if not ext.startswith("."):
            ext = "." + ext
        normalized.add(ext)
    return normalized


def _load_document(path: Path) -> LoadedDocument:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
        title = _extract_markdown_title(text, path)
        return LoadedDocument(title=title, sections=[LoadedSection(text=text, metadata={})])
    if suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="ignore")
        title = _extract_text_title(text, path)
        return LoadedDocument(title=title, sections=[LoadedSection(text=text, metadata={})])
    if suffix in {".html", ".htm"}:
        html = path.read_text(encoding="utf-8", errors="ignore")
        extractor = ReadabilityExtractor()
        article = extractor.extract_article(html)
        text = article.to_markdown(including_title=True)
        title = article.title or _fallback_title(path)
        return LoadedDocument(title=title, sections=[LoadedSection(text=text, metadata={})])
    if suffix == ".pdf":
        return _load_pdf_document(path)
    if suffix == ".docx":
        return _load_docx_document(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def _load_pdf_document(path: Path) -> LoadedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("pypdf is required to ingest PDF files") from exc

    reader = PdfReader(str(path))
    title = None
    if reader.metadata and getattr(reader.metadata, "title", None):
        title = reader.metadata.title
    if not title:
        title = _fallback_title(path)

    sections: list[LoadedSection] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if not text.strip():
            continue
        sections.append(
            LoadedSection(
                text=text,
                metadata={"page_number": index},
            )
        )
    return LoadedDocument(title=title, sections=sections)


def _load_docx_document(path: Path) -> LoadedDocument:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:
        raise RuntimeError("python-docx is required to ingest DOCX files") from exc

    doc = docx.Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n\n".join(paragraphs)
    title = _extract_text_title(text, path)
    return LoadedDocument(title=title, sections=[LoadedSection(text=text, metadata={})])


def _extract_markdown_title(text: str, path: Path) -> str:
    for line in text.splitlines():
        line = line.strip()
        if line.startswith("# "):
            return line[2:].strip()
    return _fallback_title(path)


def _extract_text_title(text: str, path: Path) -> str:
    for line in text.splitlines():
        line = line.strip()
        if line:
            return line[:80]
    return _fallback_title(path)


def _fallback_title(path: Path) -> str:
    return path.stem.replace("_", " ").strip() or path.name


def _build_ingested_document(
    file_path: Path,
    loaded: LoadedDocument,
    chunk_size: int,
    chunk_overlap: int,
    root: str | None,
) -> IngestedDocument:
    document_id = _generate_document_id(file_path)
    source_path = _make_relative_path(file_path, root)
    source_uri = f"local://{source_path}"
    file_stat = file_path.stat()
    doc_metadata = {
        "document_id": document_id,
        "source": "ingest",
        "source_path": source_path,
        "source_uri": source_uri,
        "file_name": file_path.name,
        "file_ext": file_path.suffix.lower(),
        "file_size_bytes": file_stat.st_size,
        "file_modified_at": _format_timestamp(file_stat.st_mtime),
        "ingested_at": _utc_now(),
    }

    chunks: list[IngestedChunk] = []
    chunk_index = 0
    for section_index, section in enumerate(loaded.sections):
        for chunk_text in _split_text(
            section.text,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        ):
            chunk_id = f"{document_id}_chunk_{chunk_index}"
            chunk_metadata = {
                "chunk_index": chunk_index,
                "chunk_length": len(chunk_text),
                "section_index": section_index,
            }
            chunk_metadata.update(section.metadata)
            chunks.append(
                IngestedChunk(
                    chunk_id=chunk_id,
                    content=chunk_text,
                    metadata=chunk_metadata,
                )
            )
            chunk_index += 1

    return IngestedDocument(
        document_id=document_id,
        title=loaded.title,
        url=source_uri,
        metadata=doc_metadata,
        chunks=chunks,
    )


def _split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be greater than 0")
    if chunk_overlap < 0:
        raise ValueError("chunk_overlap must be non-negative")
    if not text or not text.strip():
        return []
    if len(text) <= chunk_size:
        return [text.strip()]

    paragraphs = _split_paragraphs(text)
    chunks = []
    current = ""

    for paragraph in paragraphs:
        if len(paragraph) > chunk_size:
            if current.strip():
                chunks.append(current.strip())
                current = ""
            chunks.extend(_split_long_text(paragraph, chunk_size, chunk_overlap))
            continue
        if not current:
            current = paragraph
            continue
        if len(current) + 2 + len(paragraph) <= chunk_size:
            current = f"{current}\n\n{paragraph}"
            continue

        chunks.append(current.strip())
        if chunk_overlap > 0:
            overlap = current[-chunk_overlap:]
            current = f"{overlap}\n\n{paragraph}"
        else:
            current = paragraph

    if current.strip():
        chunks.append(current.strip())

    return chunks


def _split_long_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    chunks = []
    start = 0
    text_length = len(text)
    while start < text_length:
        end = min(start + chunk_size, text_length)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= text_length:
            break
        start = end - chunk_overlap if chunk_overlap > 0 else end
    return chunks


def _split_paragraphs(text: str) -> list[str]:
    paragraphs = []
    for block in re.split(r"\n\s*\n", text):
        block = block.strip()
        if block:
            paragraphs.append(block)
    return paragraphs


def _generate_document_id(path: Path) -> str:
    stat = path.stat()
    signature = f"{path.as_posix()}::{stat.st_size}::{stat.st_mtime}"
    digest = hashlib.md5(signature.encode("utf-8")).hexdigest()[:10]
    stem = path.stem.replace(" ", "_") or "document"
    return f"{stem}_{digest}"


def _make_relative_path(path: Path, root: str | None) -> str:
    try:
        if root:
            return str(path.resolve().relative_to(Path(root).resolve()))
        return str(path.resolve().relative_to(Path.cwd()))
    except ValueError:
        return str(path.resolve())


def _format_timestamp(value: float) -> str:
    return datetime.fromtimestamp(value, tz=timezone.utc).isoformat()


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def main() -> None:
    parser = argparse.ArgumentParser(description="Ingest documents into RAG store")
    parser.add_argument("paths", nargs="+", help="File or directory paths to ingest")
    parser.add_argument(
        "--chunk-size",
        type=int,
        default=get_int_env("RAG_CHUNK_SIZE", 1200),
        help="Chunk size in characters",
    )
    parser.add_argument(
        "--chunk-overlap",
        type=int,
        default=get_int_env("RAG_CHUNK_OVERLAP", 200),
        help="Overlap size in characters",
    )
    parser.add_argument(
        "--extensions",
        default=",".join(DEFAULT_EXTENSIONS),
        help="Comma-separated list of file extensions to ingest",
    )
    parser.add_argument(
        "--root",
        default=None,
        help="Root path used to compute relative source paths",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Parse and chunk documents without inserting into the RAG store",
    )
    parser.add_argument(
        "--report",
        default="",
        help="Optional path to write a JSON ingestion summary",
    )

    args = parser.parse_args()
    extensions = [ext.strip() for ext in args.extensions.split(",") if ext.strip()]
    summary = ingest_paths(
        paths=args.paths,
        chunk_size=args.chunk_size,
        chunk_overlap=args.chunk_overlap,
        extensions=extensions,
        root=args.root,
        dry_run=args.dry_run,
    )

    print(json.dumps(summary.to_dict(), ensure_ascii=False, indent=2))

    if args.report:
        Path(args.report).write_text(
            json.dumps(summary.to_dict(), ensure_ascii=False, indent=2),
            encoding="utf-8",
        )


if __name__ == "__main__":
    main()
